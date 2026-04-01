# -*- coding: utf-8 -*-
import sys, os
from docx import Document

QUELLE = os.path.join(
    os.path.expanduser("~"),
    "OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V",
    "Desktop", "Max30_Bul3_Orange.docx"
)

doc = Document(QUELLE)
print("=== PARAGRAPHEN ===")
for i, p in enumerate(doc.paragraphs):
    print(f"[{i}] style={p.style.name!r} text={p.text[:100]!r}")

print(f"\n=== TABELLEN: {len(doc.tables)} ===")
for ti, t in enumerate(doc.tables):
    print(f"\nTabelle {ti}: {len(t.rows)} Zeilen x {len(t.columns)} Spalten")
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text[:80].replace("\n", "\\n")
            print(f"  [{ri},{ci}]: {txt!r}")

print("\n=== SEITENGRÖSSE ===")
for s in doc.sections:
    print(f"  Breite={s.page_width.cm:.1f}cm Höhe={s.page_height.cm:.1f}cm")
    print(f"  Links={s.left_margin.cm:.1f}cm Rechts={s.right_margin.cm:.1f}cm")
