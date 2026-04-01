"""
Test-Script um die Struktur eines Stärkemeldung-Dokuments zu analysieren
"""
from docx import Document
import os

# Test-Dokument
TEST_DOC = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26\06_Stärkemeldung\03_März\Stärkemeldung 01.03.2026 - 02.03.2026.docx"

print(f"Analysiere: {os.path.basename(TEST_DOC)}\n")

doc = Document(TEST_DOC)

print("="*80)
print("ABSÄTZE (ALLE)")
print("="*80)
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i:3d}] {text[:120]}")

print("\n" + "="*80)
print("TABELLEN")
print("="*80)
for ti, table in enumerate(doc.tables):
    print(f"\nTabelle {ti}: {len(table.rows)} Zeilen x {len(table.columns)} Spalten")
    
    # Zeige die letzten paar Zeilen der Tabelle
    start_row = max(0, len(table.rows) - 5)
    for ri, row in enumerate(table.rows[start_row:], start=start_row):
        cells_text = []
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip().replace('\n', ' / ')
            if text:
                cells_text.append(f"[{ci}]: {text[:50]}")
        if cells_text:
            print(f"  Zeile {ri}: {' | '.join(cells_text)}")
