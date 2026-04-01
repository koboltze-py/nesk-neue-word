# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn

def lese_doc(pfad):
    doc = Document(pfad)
    result = {}
    for sec in doc.sections:
        result['seite'] = (round(sec.page_width.cm,2), round(sec.page_height.cm,2))
        result['raender'] = (round(sec.top_margin.cm,2), round(sec.bottom_margin.cm,2),
                             round(sec.left_margin.cm,2), round(sec.right_margin.cm,2))
    tbl = doc.tables[0]
    tr = tbl.rows[0]._tr
    trPr = tr.find(qn('w:trPr'))
    trH = trPr.find(qn('w:trHeight')) if trPr else None
    result['tbl_hoehe_cm'] = round(int(trH.get(qn('w:val')))/567, 3) if trH is not None else None
    result['col_breiten'] = [round(c.width.cm, 2) for c in tbl.rows[0].cells]

    runs = []
    def scan_cell(cell):
        for para in cell.paragraphs:
            sa = para.paragraph_format.space_after
            sb = para.paragraph_format.space_before
            for run in para.runs:
                sz = run.font.size
                pt = round(sz.pt, 1) if sz else None
                if run.text.strip():
                    runs.append((pt, run.bold, run.text.strip()[:50],
                                 round(sa.pt,1) if sa else 0,
                                 round(sb.pt,1) if sb else 0))
        for sub in cell.tables:
            for row in sub.rows:
                for sc in row.cells:
                    scan_cell(sc)

    for cell in tbl.rows[0].cells:
        scan_cell(cell)
    result['runs'] = runs
    return result

v8    = lese_doc(r'word neu 26.3\DEMO_Dashboard_v8_25032026.docx')
final = lese_doc(r'word neu 26.3\DEMO_Dashboard_final_25032026.docx')

print("=== Tabellenhoehe ===")
print("v8   :", v8['tbl_hoehe_cm'], "cm")
print("final:", final['tbl_hoehe_cm'], "cm")

print()
print("=== Spaltenbreiten ===")
print("v8   :", v8['col_breiten'])
print("final:", final['col_breiten'])

print()
print("=== Unterschiede in Runs (pt | bold | text | sa | sb) ===")
for i, (r1, r2) in enumerate(zip(v8['runs'], final['runs'])):
    if r1 != r2:
        print("  Zeile", i)
        print("    v8   :", r1)
        print("    final:", r2)

n1, n2 = len(v8['runs']), len(final['runs'])
if n1 != n2:
    print("Anz. Runs: v8=", n1, " final=", n2)
