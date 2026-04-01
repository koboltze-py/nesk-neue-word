# -*- coding: utf-8 -*-
"""
DEMO – Betreuerliste (links) NEBEN Stärkemeldung/Stats (rechts)
Quelle links : Max30_Bul3_Orange.docx (Betreuer Tag + Nacht)
Quelle rechts: F1-Stil (OzeanBlau – Stärkemeldung / Dashboard)
Layout       : A4 quer (29,7 cm × 21,0 cm)
"""

import os, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Pfade ─────────────────────────────────────────────────────────────────────
_OD  = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
LOGO = Path(__file__).resolve().parents[1] / "Daten" / "Email" / "Logo.jpg"
ZIEL = Path(__file__).parent   # word neu 26.3/
DATUM   = "25.03.2026"
STATION = "Erste-Hilfe-Station · Flughafen Köln/Bonn"
TEL     = "+49 2203 40-2323"
MAIL    = "erste-hilfe-station-flughafen@drk-koeln.de"

# ── Farbschema OzeanBlau (identisch F1) ───────────────────────────────────────
BG   = "1A3460"   # Dunkelblau – Hintergrund
DK   = "0F1F3C"   # Noch dunkler
HE   = "C8DAFF"   # Hellblau
AZ   = "00C8FF"   # Tagdienst-Akzent
AZ2  = "00A878"   # Nachtdienst-Akzent
WEIS = "FFFFFF"

# ── Demo-Daten (aus Max30_Bul3_Orange.docx ausgelesen) ────────────────────────
# Tag-Betreuer: (Name, "HH:MM–HH:MM")
TAG_BET = [
    ("Athanasiou",  "06:00–18:00"), ("Badrieh",    "06:00–18:00"),
    ("Baluch",      "06:15–18:00"), ("Bauschke",   "06:00–17:00"),
    ("Gül",         "06:00–18:30"), ("Schneider",  "06:00–18:00"),
    ("Tamer",       "06:00–18:00"), ("Üzülmez",    "06:00–18:00"),
    ("El Mojahid",  "08:00–16:00"), ("Idic",       "08:00–18:00"),
    ("Lehmann",     "08:00–16:00"), ("Peters",     "08:00–17:15"),
    ("Doubli",      "09:00–19:00"), ("Heim",       "09:00–19:00"),
    ("Loukili",     "09:00–19:00"), ("Pieper",     "09:00–19:00"),
    ("Thiebes",     "09:00–19:00"), ("Delgado",    "09:30–19:00"),
    ("Cemal",       "10:00–18:00"),
]

# Nacht-Betreuer
NACHT_BET = [
    ("Acar",        "18:00–06:00"), ("Bakkal",     "18:00–06:00"),
    ("Bedl",        "18:30–06:00"), ("Bouladhane", "18:00–06:00"),
    ("Campolo",     "18:00–06:00"), ("Dobrani",    "18:00–06:00"),
    ("Mantzas",     "18:00–06:20"), ("Oh",         "18:00–06:00"),
    ("Taute",       "18:00–06:00"), ("Chugh",      "21:00–07:00"),
    ("Hein",        "21:00–07:00"), ("Irani",      "21:00–07:00"),
    ("Isa",         "21:00–07:00"), ("Tunahan",    "21:00–07:00"),
]

# Statistik-Demo-Werte (aus der Abbildung / F1-Stärkemeldung)
BUL_ANZ   = 3            # Bulmor aktiv
EINZ      = 7            # Einsätze Schichtleiter
BET_JAHR  = 21_000       # Betreuungen 2026
BET_VORT  = 300          # Betreuungen Vortag
SCHICHTL  = "Peters"


# ═══════════════════════════════════════════════════════════════════════════════
# Hilfs-Funktionen (aus _erstelle_F1.py übernommen)
# ═══════════════════════════════════════════════════════════════════════════════
def _rgb(hx: str) -> RGBColor:
    hx = hx.lstrip("#")
    return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

def _set_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)

def _no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcB.append(b)
    tcPr.append(tcB)

def _p(cell, text, bold=False, size=9, fg="000000",
       align="left", sa=0, sb=0):
    p = cell.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                   else WD_ALIGN_PARAGRAPH.RIGHT  if align == "right"
                   else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = _rgb(fg)
    return p

def _trennlinie(cell, color_hex, oben=False):
    sep = cell.add_paragraph()
    sep.paragraph_format.space_before = Pt(3)
    sep.paragraph_format.space_after  = Pt(1)
    pPr = sep._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    edge = OxmlElement("w:top" if oben else "w:bottom")
    edge.set(qn("w:val"), "single"); edge.set(qn("w:sz"), "6")
    edge.set(qn("w:space"), "1"); edge.set(qn("w:color"), color_hex.upper())
    pBdr.append(edge); pPr.append(pBdr)

def _section_hdr(cell, text, bg_hex, fg="FFFFFF", width=Cm(14), sa=0, sb=0):
    t = cell.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.cell(0, 0); _no_border(c); _set_bg(c, bg_hex); c.width = width
    _p(c, text, bold=True, size=8, fg=fg, sa=sa, sb=sb)
    return t

def _namen_tabelle_2sp(cell, lst, w_name=Cm(3.2), w_zeit=Cm(2.0),
                       bg_a="F4F4F4", bg_b="FFFFFF", total_w=Cm(14)):
    """Namensliste in 2 Namens-Paaren nebeneinander (4 Spalten)"""
    if not lst: return
    half = (len(lst) + 1) // 2
    rows = half
    tbl = cell.add_table(rows=rows, cols=4); tbl.style = "Table Grid"
    for i, (name, zeit) in enumerate(lst):
        row = i % half; col = (i // half) * 2
        nc = tbl.cell(row, col); zc = tbl.cell(row, col + 1)
        _no_border(nc); _no_border(zc)
        bg = bg_a if row % 2 == 0 else bg_b
        _set_bg(nc, bg); _set_bg(zc, bg)
        nc.width = w_name; zc.width = w_zeit
        pn = nc.paragraphs[0]
        pn.paragraph_format.space_before = Pt(1)
        pn.paragraph_format.space_after  = Pt(1)
        rn = pn.add_run(name)
        rn.font.size = Pt(7.5); rn.font.color.rgb = _rgb("111111")
        pz = zc.paragraphs[0]
        pz.paragraph_format.space_before = Pt(1)
        pz.paragraph_format.space_after  = Pt(1)
        rz = pz.add_run(zeit)
        rz.font.size = Pt(7); rz.font.color.rgb = _rgb("667799"); rz.bold = True

def _bul_status(bul):
    if bul <= 2: return "FF3333", "KRITISCH",      "3A0000"
    if bul == 3: return "E07800", "EINGESCHRÄNKT", "3A2000"
    return               "10A050", "VOLLSTÄNDIG",   "003A18"


# ═══════════════════════════════════════════════════════════════════════════════
# HAUPT-FUNKTION: erstelle Kombinations-Dokument
# ═══════════════════════════════════════════════════════════════════════════════
def erstelle_demo():
    fc_hex, lbl_bul, fc_bg = _bul_status(BUL_ANZ)

    # ── Dokument & Seite (A4 quer) ───────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.page_width   = Cm(29.7)
        sec.page_height  = Cm(21.0)
        sec.top_margin   = Cm(0.6)
        sec.bottom_margin= Cm(0.6)
        sec.left_margin  = Cm(0.5)
        sec.right_margin = Cm(0.5)

    # ── Haupt-Tabelle: 2 Spalten ─────────────────────────────────────────────
    L_W = Cm(14.4)   # Betreuerliste (links)
    R_W = Cm(14.3)   # Stärkemeldung/Stats (rechts)

    main = doc.add_table(rows=1, cols=2)
    main.style = "Table Grid"
    lc = main.cell(0, 0)   # links  – Betreuer-Liste
    rc = main.cell(0, 1)   # rechts – Stats/Dashboard
    lc.width = L_W; rc.width = R_W
    _no_border(lc); _no_border(rc)
    _set_bg(lc, "FFFFFF"); _set_bg(rc, BG)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ══════════════════════════════════════════════════════════════════════════
    # LINKE SPALTE – Betreuerliste (aus Max30_Bul3_Orange.docx)
    # ══════════════════════════════════════════════════════════════════════════

    # Logo + Stations-Header
    try:
        if LOGO.exists():
            lp = lc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_before = Pt(3)
            lp.add_run().add_picture(str(LOGO), width=Cm(2.5))
    except Exception:
        pass

    _p(lc, "Deutsches Rotes Kreuz – Kreisverband Köln e.V.",
       bold=True, size=9, fg="1A3460", align="center", sb=2)
    _p(lc, STATION, size=7.5, fg="555555", align="center")
    _p(lc, f"BETREUER  ·  {DATUM}", bold=True, size=8.5, fg="1A3460",
       align="center", sa=3)

    _trennlinie(lc, AZ)

    # Tagdienst-Header
    _section_hdr(lc, f"BETREUER – TAGDIENST  ({len(TAG_BET)} Personen)",
                 DK, fg=AZ, width=L_W)
    _namen_tabelle_2sp(lc, TAG_BET, total_w=L_W)

    # Nachtdienst-Header
    _section_hdr(lc, f"BETREUER – NACHTDIENST  ({len(NACHT_BET)} Personen)",
                 "2A2A4A", fg=AZ2, width=L_W, sb=2)
    _namen_tabelle_2sp(lc, NACHT_BET, bg_a="EDEDFF", total_w=L_W)

    # Footer links
    sp = lc.add_paragraph(); sp.paragraph_format.space_before = Pt(6)
    ft = lc.add_table(rows=1, cols=1); ft.style = "Table Grid"
    ftc = ft.cell(0, 0); _no_border(ftc); _set_bg(ftc, DK); ftc.width = L_W
    _p(ftc, f"DRK Köln  ·  {TEL}  ·  {MAIL}",
       size=7, fg="AAAAAA", align="center", sb=2, sa=2)

    # ══════════════════════════════════════════════════════════════════════════
    # RECHTE SPALTE – Stärkemeldung / Stats (jetziger Word-Export F1-Stil)
    # ══════════════════════════════════════════════════════════════════════════

    # Logo + Org-Header
    try:
        if LOGO.exists():
            rp = rc.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rp.paragraph_format.space_before = Pt(3)
            rp.add_run().add_picture(str(LOGO), width=Cm(2.5))
    except Exception:
        pass

    _p(rc, "Deutsches Rotes Kreuz", bold=True, size=10, fg=WEIS, align="center", sb=2)
    _p(rc, "Kreisverband Köln e.V.", size=7.5, fg=HE, align="center")
    _p(rc, STATION, size=7, fg="AAAAAA", align="center")
    _p(rc, TEL,     bold=True, size=8, fg=HE, align="center", sa=2)

    _trennlinie(rc, AZ)
    _p(rc, f"📅  {DATUM}", bold=True, size=11, fg=WEIS, align="center", sb=3, sa=4)
    _trennlinie(rc, AZ)

    # Kennzahlen
    kennz = [
        ("✦ Anzahl Betreuungen (2026)", f"{BET_JAHR:,}".replace(",", "."), AZ),
        ("✦ Anzahl Betreuungen (Vortag)", str(BET_VORT),                  HE),
        ("✦ Einsätze Schichtleiter",      str(EINZ),                      "AAFFCC"),
        ("✦ Personal gesamt",             str(len(TAG_BET) + len(NACHT_BET)), HE),
    ]
    for lbl, val, vc in kennz:
        p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(3)
        r1 = p.add_run(f"{lbl}\n"); r1.font.size = Pt(7.5)
        r1.font.color.rgb = _rgb(HE)
        r2 = p.add_run(val)
        r2.bold = True; r2.font.size = Pt(18)
        r2.font.color.rgb = _rgb(vc)

    # Bulmor-Status
    _trennlinie(rc, AZ, oben=True)
    _p(rc, "BULMOR – FAHRZEUGSTATUS", bold=True, size=8.5, fg=AZ,
       align="center", sb=2)

    st = rc.add_table(rows=1, cols=1); st.style = "Table Grid"
    stc = st.cell(0, 0); _no_border(stc); _set_bg(stc, fc_bg); stc.width = R_W
    p_st = stc.paragraphs[0]; p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_st = p_st.add_run(f"Gesamt:  {BUL_ANZ}/5  ·  {lbl_bul}")
    r_st.bold = True; r_st.font.size = Pt(11)
    r_st.font.color.rgb = _rgb(fc_hex)

    # Fahrzeug-Symbole B1–B5 als Tabelle
    vtbl = rc.add_table(rows=2, cols=5); vtbl.style = "Table Grid"
    for i in range(5):
        aktiv = (i + 1) <= BUL_ANZ
        c_top = vtbl.cell(0, i)
        c_bot = vtbl.cell(1, i)
        bg_top = fc_hex if aktiv else "CC4444"
        bg_bot = "EEF6FF" if aktiv else "F8E0E0"
        _no_border(c_top); _no_border(c_bot)
        _set_bg(c_top, bg_top); _set_bg(c_bot, bg_bot)
        p_t = c_top.paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_t.paragraph_format.space_before = Pt(2)
        p_t.paragraph_format.space_after  = Pt(2)
        rt = p_t.add_run(f"●\nB{i+1}")
        rt.bold = True; rt.font.size = Pt(8)
        rt.font.color.rgb = _rgb(WEIS)
        p_b = c_bot.paragraphs[0]
        p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_b.paragraph_format.space_before = Pt(1)
        p_b.paragraph_format.space_after  = Pt(1)
        rb = p_b.add_run("Dienst" if aktiv else "Aus")
        rb.font.size = Pt(7.5)
        rb.font.color.rgb = _rgb(fc_hex if aktiv else "993333")

    # Schichtleiter
    _trennlinie(rc, AZ, oben=True)
    _p(rc, "SCHICHTLEITER", bold=True, size=8, fg=AZ, align="center", sb=2)
    _p(rc, SCHICHTL,        bold=True, size=12, fg=WEIS, align="center", sa=2)

    # Footer rechts
    sp2 = rc.add_paragraph(); sp2.paragraph_format.space_before = Pt(6)
    ft2 = rc.add_table(rows=1, cols=1); ft2.style = "Table Grid"
    ft2c = ft2.cell(0, 0); _no_border(ft2c); _set_bg(ft2c, DK); ft2c.width = R_W
    _p(ft2c, f"DRK Köln  ·  {TEL}  ·  {MAIL}",
       size=7, fg="AAAAAA", align="center", sb=2, sa=2)

    # ── Speichern ─────────────────────────────────────────────────────────────
    out = ZIEL / "DEMO_Betreuer_links_Staerkemeldung_rechts_25032026.docx"
    doc.save(str(out))
    print(f"[OK] Demo gespeichert: {out}")
    return str(out)


if __name__ == "__main__":
    erstelle_demo()
